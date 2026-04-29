import docx
import sys

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also read tables
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip().replace('\n', ' '))
                full_text.append(" | ".join(row_data))

        print(f"--- Content of {file_path} ---")
        print("\n".join(full_text))
        print(f"--- End of {file_path} ---")
    except Exception as e:
        print(f"Error reading {file_path}: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        for f in sys.argv[1:]:
            read_docx(f)
